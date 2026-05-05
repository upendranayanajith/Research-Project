"""
Generate External Supervisor Feedback Form as .docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CHECKBOX = "☐"  # ☐

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def make_table(rows, cols, header=None, widths=None, header_shade='2E5984'):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Light Grid Accent 1'
    t.autofit = True
    if widths:
        for i, w in enumerate(widths):
            for cell in t.columns[i].cells:
                cell.width = w
    if header:
        hdr = t.rows[0].cells
        for i, txt in enumerate(header):
            hdr[i].text = ''
            p = hdr[i].paragraphs[0]
            r = p.add_run(txt)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade_cell(hdr[i], header_shade)
    return t


# ===========================================================================
# TITLE
# ===========================================================================
title = doc.add_heading('External Supervisor Feedback Form', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Project Title: ')
r.bold = True
p.add_run('High-Accuracy Reading Pipeline for Industrial Analog Gauges with Vision-Language Shadow Filtering')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Intended Application Domain: ')
r.bold = True
p.add_run('Garment Industry Machinery Integration')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Date of Review: ')
r.bold = True
p.add_run('____ / ____ / 20____')

doc.add_paragraph()

# ===========================================================================
# SECTION A
# ===========================================================================
add_heading('Section A — Supervisor Information', level=1)

t = make_table(rows=9, cols=2, header=['Field', 'Detail'])
fields_a = [
    'Full Name',
    'Designation / Position',
    'Company / Organization',
    'Department / Division',
    'Years of Industry Experience',
    'Email',
    'Contact Number',
    'Relationship to Project',
]
for i, f in enumerate(fields_a, start=1):
    t.rows[i].cells[0].text = f
    if f == 'Relationship to Project':
        t.rows[i].cells[1].text = (
            f'{CHECKBOX} Project Sponsor   {CHECKBOX} Technical Advisor   '
            f'{CHECKBOX} End-User Representative   {CHECKBOX} Other: ________'
        )
    else:
        t.rows[i].cells[1].text = ''

doc.add_paragraph()

# ===========================================================================
# SECTION B
# ===========================================================================
add_heading('Section B — Project Requirements Alignment', level=1)
add_para('Please rate how well the delivered system aligns with the original requirements you specified.', italic=True)

b_items = [
    ('B1', 'The system addresses the original problem you raised for the garment industry'),
    ('B2', 'The scope of the gauge / instrument types covered matches what your machinery uses'),
    ('B3', 'The reading accuracy reported is sufficient for your operational tolerances'),
    ('B4', 'The latency (~186 ms per reading) is acceptable for your line-speed / sampling needs'),
    ('B5', 'The shadow-handling capability matches the lighting conditions in your facilities'),
]
t = make_table(rows=len(b_items)+1, cols=7,
               header=['#', 'Criterion', '1 (Poor)', '2', '3', '4', '5 (Excellent)'])
for i, (num, crit) in enumerate(b_items, start=1):
    t.rows[i].cells[0].text = num
    t.rows[i].cells[1].text = crit
    for j in range(2, 7):
        t.rows[i].cells[j].text = CHECKBOX
        t.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_para('Comments on requirements alignment:', bold=True)
for _ in range(3):
    doc.add_paragraph('_' * 95)

doc.add_paragraph()

# ===========================================================================
# SECTION C
# ===========================================================================
add_heading('Section C — Technical Quality (1–5 scale)', level=1)
c_items = [
    ('C1', 'Soundness of the four-stage pipeline design'),
    ('C2', 'Quality of the shadow-filtering approach (CLIP analysis-by-synthesis)'),
    ('C3', 'Robustness under varying illumination (LED, fluorescent, daylight)'),
    ('C4', 'Calibration / uncertainty reporting (ECE, confidence taxonomy)'),
    ('C5', 'Edge-deployment readiness (Jetson AGX Orin INT8)'),
    ('C6', 'Failure-mode analysis depth'),
]
t = make_table(rows=len(c_items)+1, cols=4,
               header=['#', 'Criterion', 'Score (1–5)', 'Comments'])
for i, (num, crit) in enumerate(c_items, start=1):
    t.rows[i].cells[0].text = num
    t.rows[i].cells[1].text = crit
    t.rows[i].cells[2].text = f'{CHECKBOX}1  {CHECKBOX}2  {CHECKBOX}3  {CHECKBOX}4  {CHECKBOX}5'
    t.rows[i].cells[3].text = ''

doc.add_paragraph()

# ===========================================================================
# SECTION D
# ===========================================================================
add_heading('Section D — Garment Industry Integration Suitability', level=1)
add_para('Please assess fitness for deployment on your specific machinery.', italic=True)

add_para('D1. Which garment-industry machines are you intending this system to monitor? (tick all that apply)', bold=True)
machines = [
    'Steam pressure gauges (boilers, irons, steam tunnels)',
    'Compressed-air pressure gauges (pneumatic cutters, blowers)',
    'Temperature dials (dyeing vats, drying chambers, heat-set tunnels)',
    'Speedometers / RPM gauges (sewing machines, spinning units)',
    'Flow meters (water/dye/chemical lines)',
    'Tension gauges (knitting / weaving machinery)',
    'Other: ____________________________________________',
]
for m in machines:
    p = doc.add_paragraph(f'{CHECKBOX}  {m}', style='List Bullet')

add_para('D2. Operating environment factors at your facility (tick all relevant)', bold=True)
factors = [
    'High humidity (>70%)',
    'Lint / fibre dust accumulation on dial faces',
    'Vibration from machinery',
    'Steam / condensation on gauge glass',
    'Mixed natural + artificial lighting',
    '24/7 continuous operation',
    'Multiple gauges per machine (specify count): ______',
]
for f in factors:
    doc.add_paragraph(f'{CHECKBOX}  {f}', style='List Bullet')

add_para('D3. How well does the demonstrated pipeline handle these conditions?', bold=True)
doc.add_paragraph(
    f'{CHECKBOX} Fully    {CHECKBOX} Mostly    {CHECKBOX} Partially    '
    f'{CHECKBOX} Not adequately    {CHECKBOX} Not yet evaluated'
)

add_para('Specific concerns about garment-environment robustness:', bold=True)
for _ in range(3):
    doc.add_paragraph('_' * 95)

add_para('D4. Integration with existing factory systems', bold=True)
integ = [
    'Compatibility with existing PLCs / SCADA',
    'MES / ERP data-export format',
    'Real-time alarm / threshold triggering',
    'Operator dashboard / HMI usability',
    'OPC-UA / Modbus / MQTT support',
]
t = make_table(rows=len(integ)+1, cols=3,
               header=['Requirement', 'Met?', 'Comments'])
for i, req in enumerate(integ, start=1):
    t.rows[i].cells[0].text = req
    t.rows[i].cells[1].text = f'{CHECKBOX} Yes  {CHECKBOX} Partial  {CHECKBOX} No  {CHECKBOX} N/A'
    t.rows[i].cells[2].text = ''

doc.add_paragraph()

# ===========================================================================
# SECTION E
# ===========================================================================
add_heading('Section E — Deployment Readiness', level=1)
e_items = [
    ('E1', 'Hardware specification matches available factory IT/edge infrastructure'),
    ('E2', 'Camera mounting and field-of-view requirements are practical for our machines'),
    ('E3', 'Power and network requirements are achievable on the production floor'),
    ('E4', 'IP-65 / dust-resistant enclosure specified for garment-floor conditions'),
    ('E5', 'Documentation is sufficient for our maintenance team to operate'),
    ('E6', 'Training material for line operators is adequate'),
]
t = make_table(rows=len(e_items)+1, cols=3,
               header=['#', 'Item', 'Status'])
for i, (num, item) in enumerate(e_items, start=1):
    t.rows[i].cells[0].text = num
    t.rows[i].cells[1].text = item
    t.rows[i].cells[2].text = f'{CHECKBOX} Ready   {CHECKBOX} Needs work   {CHECKBOX} Not ready'

doc.add_paragraph()

# ===========================================================================
# SECTION F
# ===========================================================================
add_heading('Section F — Business / Operational Impact', level=1)

add_para('F1. Expected benefit to your garment operation (tick all)', bold=True)
benefits = [
    'Reduced manual gauge-reading rounds',
    'Earlier detection of pressure / temperature deviations',
    'Reduced fabric defects from out-of-spec parameters',
    'Energy / steam savings via tighter monitoring',
    'Compliance / audit trail',
    'Insurance / safety reporting',
    'Other: ____________________________________________',
]
for b in benefits:
    doc.add_paragraph(f'{CHECKBOX}  {b}', style='List Bullet')

add_para('F2. Estimated number of machines / gauges for initial pilot deployment: ______', bold=True)
add_para('F3. Estimated scale-up if pilot succeeds (number of gauges): ______', bold=True)
add_para('F4. Anticipated ROI horizon:', bold=True)
doc.add_paragraph(
    f'{CHECKBOX} <6 months    {CHECKBOX} 6–12 months    '
    f'{CHECKBOX} 1–2 years    {CHECKBOX} >2 years'
)

doc.add_paragraph()

# ===========================================================================
# SECTION G
# ===========================================================================
add_heading('Section G — Strengths Observed', level=1)
add_para('Please list aspects that exceeded expectations or that are particularly well-suited to garment-industry deployment:', italic=True)
for n in range(1, 5):
    doc.add_paragraph(f'{n}. ' + '_' * 92)

doc.add_paragraph()

# ===========================================================================
# SECTION H
# ===========================================================================
add_heading('Section H — Gaps / Concerns', level=1)
add_para('Please list aspects that need improvement before factory roll-out:', italic=True)
t = make_table(rows=5, cols=3, header=['Concern', 'Severity', 'Suggested Action'])
for i in range(1, 5):
    t.rows[i].cells[0].text = ''
    t.rows[i].cells[1].text = f'{CHECKBOX} Low  {CHECKBOX} Med  {CHECKBOX} High'
    t.rows[i].cells[2].text = ''

doc.add_paragraph()

# ===========================================================================
# SECTION I
# ===========================================================================
add_heading('Section I — Required Modifications Before Pilot', level=1)
add_para('What changes must be made before the system can be trialled on your factory floor?', italic=True)
for n in range(1, 5):
    doc.add_paragraph(f'{n}. ' + '_' * 92)

doc.add_paragraph()

# ===========================================================================
# SECTION J
# ===========================================================================
add_heading('Section J — Open Feedback', level=1)

add_para('J1. What did the team do exceptionally well?', bold=True)
for _ in range(3):
    doc.add_paragraph('_' * 95)

add_para('J2. What was the weakest aspect of the project as delivered?', bold=True)
for _ in range(3):
    doc.add_paragraph('_' * 95)

add_para('J3. Suggestions for further research / future phases (e.g., needle-thread tension monitoring, fabric-defect vision, predictive maintenance):', bold=True)
for _ in range(3):
    doc.add_paragraph('_' * 95)

add_para('J4. Would you recommend continuing this collaboration into a pilot deployment?', bold=True)
doc.add_paragraph(
    f'{CHECKBOX} Yes, immediately    {CHECKBOX} Yes, after revisions    '
    f'{CHECKBOX} Undecided    {CHECKBOX} No'
)

doc.add_paragraph()

# ===========================================================================
# SECTION K
# ===========================================================================
add_heading('Section K — Overall Assessment', level=1)

add_para('Overall Project Rating:', bold=True)
rating_line = '   '.join(f'{CHECKBOX} {n}' for n in range(1, 11))
doc.add_paragraph(rating_line)

add_para('Final Recommendation:', bold=True)
recs = [
    'Accept as-is for pilot deployment',
    'Accept with minor modifications (listed above)',
    'Accept with major modifications (listed above)',
    'Re-evaluate after substantial rework',
    'Not suitable for our use case',
]
for rec in recs:
    doc.add_paragraph(f'{CHECKBOX}  {rec}', style='List Bullet')

add_para('Justification:', bold=True)
for _ in range(4):
    doc.add_paragraph('_' * 95)

doc.add_paragraph()

# ===========================================================================
# SECTION L
# ===========================================================================
add_heading('Section L — Declaration & Signature', level=1)
add_para(
    'I confirm that I have reviewed the project deliverables (research paper, '
    'demonstration, and supporting documentation) and that the feedback above '
    'represents my honest professional assessment regarding suitability for '
    'integration into garment industry machinery.', italic=True)

doc.add_paragraph()

t = make_table(rows=5, cols=2)
t.rows[0].cells[0].text = 'Signature:'
t.rows[0].cells[1].text = '_____________________________'
t.rows[1].cells[0].text = 'Name (printed):'
t.rows[1].cells[1].text = '_____________________________'
t.rows[2].cells[0].text = 'Designation:'
t.rows[2].cells[1].text = '_____________________________'
t.rows[3].cells[0].text = 'Company stamp / seal:'
t.rows[3].cells[1].text = ''
t.rows[4].cells[0].text = 'Date:'
t.rows[4].cells[1].text = '____ / ____ / 20____'

for row in t.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                if cell == row.cells[0]:
                    r.bold = True

doc.add_paragraph()

# Footer note
p = add_para(
    'Please return the completed form to the project team within 14 days of '
    'receipt. For clarifications, contact the academic supervisor at the '
    'institution.', italic=True, size=9)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===========================================================================
# SAVE
# ===========================================================================
out_path = r'D:\Y4S1\Research 4\Clock_Time_Research\Research-Project\External_Supervisor_Feedback_Form.docx'
doc.save(out_path)
print(f'SUCCESS: {out_path}')
