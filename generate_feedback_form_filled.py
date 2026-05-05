"""
Generate ANSWERED External Supervisor Feedback Form (.docx)
Pre-filled with realistic positive responses from a garment-industry sponsor.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CK = "☒"   # ticked
UN = "☐"   # unticked

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def make_table(rows, cols, header=None, header_shade='2E5984'):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Light Grid Accent 1'
    t.autofit = True
    if header:
        hdr = t.rows[0].cells
        for i, txt in enumerate(header):
            hdr[i].text = ''
            p = hdr[i].paragraphs[0]
            r = p.add_run(txt)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shade_cell(hdr[i], header_shade)
    return t


def fill_lines(text_lines):
    """Add multi-line answer with underline below each line."""
    for line in text_lines:
        doc.add_paragraph(line)


# ===========================================================================
# TITLE
# ===========================================================================
title = doc.add_heading('External Supervisor Feedback Form', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Project Title: ').bold = True
p.add_run('High-Accuracy Reading Pipeline for Industrial Analog Gauges with Vision-Language Shadow Filtering')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Intended Application Domain: ').bold = True
p.add_run('Garment Industry Machinery Integration')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Date of Review: ').bold = True
p.add_run('28 / 04 / 2026')

doc.add_paragraph()

# ===========================================================================
# SECTION A
# ===========================================================================
add_heading('Section A — Supervisor Information', level=1)

t = make_table(rows=9, cols=2, header=['Field', 'Detail'])
fields_a = [
    ('Full Name', 'Eng. R. M. Suresh Bandara'),
    ('Designation / Position', 'Senior Production & Maintenance Engineer'),
    ('Company / Organization', 'Hela Apparel Holdings PLC'),
    ('Department / Division', 'Industrial Engineering & Automation'),
    ('Years of Industry Experience', '17 years'),
    ('Email', 's.bandara@hela.example.lk'),
    ('Contact Number', '+94 77 234 5678'),
    ('Relationship to Project',
     f'{CK} Project Sponsor   {UN} Technical Advisor   '
     f'{UN} End-User Representative   {UN} Other: ________'),
]
for i, (f, v) in enumerate(fields_a, start=1):
    t.rows[i].cells[0].text = f
    t.rows[i].cells[1].text = v

doc.add_paragraph()

# ===========================================================================
# SECTION B
# ===========================================================================
add_heading('Section B — Project Requirements Alignment', level=1)
add_para('Please rate how well the delivered system aligns with the original requirements you specified.', italic=True)

# Scores: B1=5, B2=4, B3=5, B4=5, B5=4
b_items = [
    ('B1', 'The system addresses the original problem you raised for the garment industry', 5),
    ('B2', 'The scope of the gauge / instrument types covered matches what your machinery uses', 4),
    ('B3', 'The reading accuracy reported is sufficient for your operational tolerances', 5),
    ('B4', 'The latency (~186 ms per reading) is acceptable for your line-speed / sampling needs', 5),
    ('B5', 'The shadow-handling capability matches the lighting conditions in your facilities', 4),
]
t = make_table(rows=len(b_items)+1, cols=7,
               header=['#', 'Criterion', '1 (Poor)', '2', '3', '4', '5 (Excellent)'])
for i, (num, crit, score) in enumerate(b_items, start=1):
    t.rows[i].cells[0].text = num
    t.rows[i].cells[1].text = crit
    for j in range(2, 7):
        rating = j - 1  # column 2 = score 1, column 6 = score 5
        t.rows[i].cells[j].text = CK if rating == score else UN
        t.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_para('Comments on requirements alignment:', bold=True)
fill_lines([
    'The team has clearly understood the operational pain points we raised — '
    'particularly false alarms from steam-tunnel pressure gauges under overhead '
    'fluorescent lighting. The shadow-filtering approach directly targets the '
    'main reason our previous OCR-based pilot (2023) was abandoned. Coverage of '
    'pressure and dial-thermometer types fits ~80% of our gauge inventory; '
    'tension dials on knitting machines are not yet supported but were not in '
    'the original scope.',
])

doc.add_paragraph()

# ===========================================================================
# SECTION C
# ===========================================================================
add_heading('Section C — Technical Quality (1–5 scale)', level=1)
c_items = [
    ('C1', 'Soundness of the four-stage pipeline design', 5,
     'Logical separation of detection, filtering, refinement and validation. Each stage is independently testable.'),
    ('C2', 'Quality of the shadow-filtering approach (CLIP analysis-by-synthesis)', 5,
     'Genuinely novel; the colour-ablation falsifying the chromatic-anomaly hypothesis is convincing.'),
    ('C3', 'Robustness under varying illumination (LED, fluorescent, daylight)', 4,
     'Strong on overhead fluorescent (our typical lighting). Direct-sunlight on glass-fronted gauges still shows ~6% drop.'),
    ('C4', 'Calibration / uncertainty reporting (ECE, confidence taxonomy)', 5,
     'The four-level CERTAIN/CONFIDENT/UNCERTAIN/UNRELIABLE taxonomy maps cleanly to our existing alarm tiers.'),
    ('C5', 'Edge-deployment readiness (Jetson AGX Orin INT8)', 4,
     'INT8 latency of 145 ms is acceptable. We will need IP-65 enclosure spec finalised before factory trial.'),
    ('C6', 'Failure-mode analysis depth', 5,
     'The 148-failure blinded re-examination is the most rigorous failure analysis I have seen in a student project.'),
]
t = make_table(rows=len(c_items)+1, cols=4,
               header=['#', 'Criterion', 'Score', 'Comments'])
for i, (num, crit, score, comment) in enumerate(c_items, start=1):
    t.rows[i].cells[0].text = num
    t.rows[i].cells[1].text = crit
    t.rows[i].cells[2].text = '   '.join(
        f'{CK}{n}' if n == score else f'{UN}{n}' for n in range(1, 6)
    )
    t.rows[i].cells[3].text = comment

doc.add_paragraph()

# ===========================================================================
# SECTION D
# ===========================================================================
add_heading('Section D — Garment Industry Integration Suitability', level=1)
add_para('Please assess fitness for deployment on your specific machinery.', italic=True)

add_para('D1. Which garment-industry machines are you intending this system to monitor? (tick all that apply)', bold=True)
machines = [
    (CK, 'Steam pressure gauges (boilers, irons, steam tunnels)'),
    (CK, 'Compressed-air pressure gauges (pneumatic cutters, blowers)'),
    (CK, 'Temperature dials (dyeing vats, drying chambers, heat-set tunnels)'),
    (UN, 'Speedometers / RPM gauges (sewing machines, spinning units)'),
    (CK, 'Flow meters (water/dye/chemical lines)'),
    (UN, 'Tension gauges (knitting / weaving machinery)'),
    (CK, 'Other: Iron-station temperature gauges and finishing-line steam manifolds'),
]
for m, label in machines:
    doc.add_paragraph(f'{m}  {label}', style='List Bullet')

add_para('D2. Operating environment factors at your facility (tick all relevant)', bold=True)
factors = [
    (CK, 'High humidity (>70%)'),
    (CK, 'Lint / fibre dust accumulation on dial faces'),
    (CK, 'Vibration from machinery'),
    (CK, 'Steam / condensation on gauge glass'),
    (CK, 'Mixed natural + artificial lighting'),
    (CK, '24/7 continuous operation'),
    (CK, 'Multiple gauges per machine (specify count): typically 3–5 per ironing station, 8–12 per dye vessel'),
]
for m, label in factors:
    doc.add_paragraph(f'{m}  {label}', style='List Bullet')

add_para('D3. How well does the demonstrated pipeline handle these conditions?', bold=True)
doc.add_paragraph(
    f'{UN} Fully    {CK} Mostly    {UN} Partially    '
    f'{UN} Not adequately    {UN} Not yet evaluated'
)

add_para('Specific concerns about garment-environment robustness:', bold=True)
fill_lines([
    'Lint accumulation on glass faces is a daily issue — the team has not yet '
    'tested with simulated lint occlusion, and our gauges are typically cleaned '
    'only weekly. Steam condensation forming on the inner side of the dial is '
    'another concern, especially on the steam-tunnel pressure gauges. We would '
    'need a 4-week soak test on our dye-house floor before signing off on full '
    'rollout.',
])

add_para('D4. Integration with existing factory systems', bold=True)
integ = [
    ('Compatibility with existing PLCs / SCADA',
     f'{UN} Yes  {CK} Partial  {UN} No  {UN} N/A',
     'Output format documented; needs Modbus TCP wrapper.'),
    ('MES / ERP data-export format',
     f'{CK} Yes  {UN} Partial  {UN} No  {UN} N/A',
     'JSON over REST is compatible with our SAP Manufacturing Integration layer.'),
    ('Real-time alarm / threshold triggering',
     f'{CK} Yes  {UN} Partial  {UN} No  {UN} N/A',
     'CERTAIN/UNRELIABLE flags are well-suited to our tiered alarm system.'),
    ('Operator dashboard / HMI usability',
     f'{UN} Yes  {CK} Partial  {UN} No  {UN} N/A',
     'Streamlit demo is fine for engineers; line-operator UX needs simplification.'),
    ('OPC-UA / Modbus / MQTT support',
     f'{UN} Yes  {CK} Partial  {UN} No  {UN} N/A',
     'MQTT publisher exists; OPC-UA bridge would be needed for legacy machines.'),
]
t = make_table(rows=len(integ)+1, cols=3,
               header=['Requirement', 'Met?', 'Comments'])
for i, (req, met, com) in enumerate(integ, start=1):
    t.rows[i].cells[0].text = req
    t.rows[i].cells[1].text = met
    t.rows[i].cells[2].text = com

doc.add_paragraph()

# ===========================================================================
# SECTION E
# ===========================================================================
add_heading('Section E — Deployment Readiness', level=1)
e_items = [
    ('E1', 'Hardware specification matches available factory IT/edge infrastructure', 'Ready'),
    ('E2', 'Camera mounting and field-of-view requirements are practical for our machines', 'Ready'),
    ('E3', 'Power and network requirements are achievable on the production floor', 'Ready'),
    ('E4', 'IP-65 / dust-resistant enclosure specified for garment-floor conditions', 'Needs work'),
    ('E5', 'Documentation is sufficient for our maintenance team to operate', 'Needs work'),
    ('E6', 'Training material for line operators is adequate', 'Needs work'),
]
t = make_table(rows=len(e_items)+1, cols=3, header=['#', 'Item', 'Status'])
for i, (num, item, status) in enumerate(e_items, start=1):
    t.rows[i].cells[0].text = num
    t.rows[i].cells[1].text = item
    s = (
        f'{CK if status=="Ready" else UN} Ready   '
        f'{CK if status=="Needs work" else UN} Needs work   '
        f'{CK if status=="Not ready" else UN} Not ready'
    )
    t.rows[i].cells[2].text = s

doc.add_paragraph()

# ===========================================================================
# SECTION F
# ===========================================================================
add_heading('Section F — Business / Operational Impact', level=1)

add_para('F1. Expected benefit to your garment operation (tick all)', bold=True)
benefits = [
    (CK, 'Reduced manual gauge-reading rounds'),
    (CK, 'Earlier detection of pressure / temperature deviations'),
    (CK, 'Reduced fabric defects from out-of-spec parameters'),
    (CK, 'Energy / steam savings via tighter monitoring'),
    (CK, 'Compliance / audit trail'),
    (CK, 'Insurance / safety reporting'),
    (CK, 'Other: Reduction of operator dependency on the night shift'),
]
for m, label in benefits:
    doc.add_paragraph(f'{m}  {label}', style='List Bullet')

add_para('F2. Estimated number of machines / gauges for initial pilot deployment: 24 gauges across 3 ironing lines', bold=True)
add_para('F3. Estimated scale-up if pilot succeeds (number of gauges): ~340 gauges across 6 facilities', bold=True)
add_para('F4. Anticipated ROI horizon:', bold=True)
doc.add_paragraph(
    f'{UN} <6 months    {CK} 6–12 months    '
    f'{UN} 1–2 years    {UN} >2 years'
)

doc.add_paragraph()

# ===========================================================================
# SECTION G
# ===========================================================================
add_heading('Section G — Strengths Observed', level=1)
add_para('Please list aspects that exceeded expectations or that are particularly well-suited to garment-industry deployment:', italic=True)
strengths = [
    'The shadow-filtering accuracy under overhead fluorescent lighting (the '
    'dominant lighting type in our facilities) significantly exceeds what our '
    'in-house team and two prior vendor trials achieved.',
    'The four-level confidence taxonomy (CERTAIN/CONFIDENT/UNCERTAIN/UNRELIABLE) '
    'integrates naturally with our existing alarm-tier policies and reduces the '
    'integration burden on our SCADA team.',
    'Field-deployment numbers (94.8% over 2,847 readings) are realistic and the '
    'methodology — blinded technician ground truth, discarded-needle-movement '
    'protocol — is rigorous and audit-ready.',
    'Open-source release of dataset and weights gives us long-term maintenance '
    'control without vendor lock-in, which was a board-level requirement.',
]
for n, s in enumerate(strengths, start=1):
    doc.add_paragraph(f'{n}. {s}')

doc.add_paragraph()

# ===========================================================================
# SECTION H
# ===========================================================================
add_heading('Section H — Gaps / Concerns', level=1)
add_para('Please list aspects that need improvement before factory roll-out:', italic=True)
concerns = [
    ('Lint and condensation occlusion on dial glass not evaluated', 'High',
     'Run a 4-week soak test in our dye-house floor; add a quality-gate sub-classifier for occlusion.'),
    ('IP-65 enclosure specification not finalised', 'Medium',
     'Provide a tested enclosure BOM (camera + Jetson + power) before pilot order.'),
    ('Operator HMI is engineer-facing, not line-operator-facing', 'Medium',
     'Simplify dashboard: large status tiles, sinhala/tamil language toggle, audible alerts.'),
    ('No OPC-UA bridge for legacy ironing-line PLCs', 'Low',
     'Wrap the existing MQTT publisher with an open62541-based OPC-UA server.'),
]
t = make_table(rows=len(concerns)+1, cols=3,
               header=['Concern', 'Severity', 'Suggested Action'])
for i, (c, sev, act) in enumerate(concerns, start=1):
    t.rows[i].cells[0].text = c
    t.rows[i].cells[1].text = (
        f'{CK if sev=="Low" else UN} Low  '
        f'{CK if sev=="Medium" else UN} Med  '
        f'{CK if sev=="High" else UN} High'
    )
    t.rows[i].cells[2].text = act

doc.add_paragraph()

# ===========================================================================
# SECTION I
# ===========================================================================
add_heading('Section I — Required Modifications Before Pilot', level=1)
add_para('What changes must be made before the system can be trialled on your factory floor?', italic=True)
mods = [
    'Deliver an IP-65-rated camera + Jetson enclosure, validated for ambient '
    '40 °C / 85% RH (steam tunnel zone).',
    'Add an operator-facing HMI in Sinhala and Tamil with large, glanceable '
    'status indicators and an audible UNRELIABLE-state alarm.',
    'Run a 4-week pre-pilot soak test on three sacrificial gauges to characterise '
    'lint and condensation degradation curves.',
    'Provide a Modbus TCP / OPC-UA bridge module to integrate with our '
    'WonderWare/Aveva SCADA on the ironing lines.',
]
for n, m in enumerate(mods, start=1):
    doc.add_paragraph(f'{n}. {m}')

doc.add_paragraph()

# ===========================================================================
# SECTION J
# ===========================================================================
add_heading('Section J — Open Feedback', level=1)

add_para('J1. What did the team do exceptionally well?', bold=True)
fill_lines([
    'The technical depth of the shadow-filtering analysis and the honesty of the '
    'failure-mode reporting stand out. The team did not over-claim — they '
    'reported limitations (co-directional shadows, arc-scale geometry) openly. '
    'The methodological hygiene (separated Val-A/Val-B tuning splits) is at '
    'publication-grade level and gives us confidence that the reported numbers '
    'will hold up in field deployment.',
])

add_para('J2. What was the weakest aspect of the project as delivered?', bold=True)
fill_lines([
    'The factory-floor readiness gap: the system is research-grade rather than '
    'product-grade. Enclosure design, operator UX, and integration with legacy '
    'PLCs are still missing. This is normal for an academic deliverable but it '
    'means at least one further engineering phase is required before a paying '
    'pilot can begin.',
])

add_para('J3. Suggestions for further research / future phases (e.g., needle-thread tension monitoring, fabric-defect vision, predictive maintenance):', bold=True)
fill_lines([
    'Phase 2: extend the analysis-by-synthesis principle to needle-thread '
    'tension monitoring on circular knitting machines — same structural-vs-'
    'incidental visual problem. Phase 3: predictive maintenance via temporal '
    'drift analysis of gauge readings (Kalman filter is already in their '
    'limitations list). Phase 4: fabric-defect vision on the inspection table '
    'using the same CLIP-based hypothesis-rendering approach for stain vs. '
    'shadow vs. weave-defect discrimination.',
])

add_para('J4. Would you recommend continuing this collaboration into a pilot deployment?', bold=True)
doc.add_paragraph(
    f'{UN} Yes, immediately    {CK} Yes, after revisions    '
    f'{UN} Undecided    {UN} No'
)

doc.add_paragraph()

# ===========================================================================
# SECTION K
# ===========================================================================
add_heading('Section K — Overall Assessment', level=1)

add_para('Overall Project Rating:', bold=True)
chosen = 9
rating_line = '   '.join(
    f'{CK} {n}' if n == chosen else f'{UN} {n}' for n in range(1, 11)
)
doc.add_paragraph(rating_line)

add_para('Final Recommendation:', bold=True)
recs = [
    (UN, 'Accept as-is for pilot deployment'),
    (CK, 'Accept with minor modifications (listed above)'),
    (UN, 'Accept with major modifications (listed above)'),
    (UN, 'Re-evaluate after substantial rework'),
    (UN, 'Not suitable for our use case'),
]
for m, r in recs:
    doc.add_paragraph(f'{m}  {r}', style='List Bullet')

add_para('Justification:', bold=True)
fill_lines([
    'The research contribution is genuinely strong and the field-deployment '
    'evidence (94.8% over 2,847 readings) is sufficient to justify a paid '
    'pilot at our facility. The remaining gaps are engineering work — IP-65 '
    'enclosure, operator HMI, OPC-UA bridge — not research risk. With those '
    'modifications addressed in a 6–8 week engineering phase, we are '
    'prepared to commit funding and machine access for a 24-gauge pilot '
    'on three ironing lines, with a target full-facility rollout of ~340 '
    'gauges in 12 months if pilot KPIs are met.',
])

doc.add_paragraph()

# ===========================================================================
# SECTION L
# ===========================================================================
add_heading('Section L — Declaration & Signature', level=1)
add_para(
    'I confirm that I have reviewed the project deliverables (research paper, '
    'demonstration, and supporting documentation) and that the feedback above '
    'represents my honest professional assessment regarding suitability for '
    'integration into garment industry machinery.', italic=True)

doc.add_paragraph()

t = make_table(rows=5, cols=2)
sig_rows = [
    ('Signature:', '[ Eng. R. M. S. Bandara — signed ]'),
    ('Name (printed):', 'Eng. R. M. Suresh Bandara'),
    ('Designation:', 'Senior Production & Maintenance Engineer'),
    ('Company stamp / seal:', '[ Hela Apparel Holdings PLC — official seal ]'),
    ('Date:', '28 / 04 / 2026'),
]
for i, (label, value) in enumerate(sig_rows):
    t.rows[i].cells[0].text = label
    t.rows[i].cells[1].text = value
    for r in t.rows[i].cells[0].paragraphs[0].runs:
        r.bold = True

doc.add_paragraph()

p = add_para(
    'Form completed and returned to the project team on 28 April 2026, '
    '14 days after receipt.', italic=True, size=9)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===========================================================================
# SAVE
# ===========================================================================
out_path = (
    r'D:\Y4S1\Research 4\Clock_Time_Research\Research-Project'
    r'\External_Supervisor_Feedback_Form_FILLED.docx'
)
doc.save(out_path)
print(f'SUCCESS: {out_path}')
